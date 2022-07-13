'''
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt

# 创建文档
document = Document()
style = document.styles['Normal']
# 标题
t0 = document.add_heading('标题0', 0)
# 居中
t0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('标题1', 1)
# 首行缩进两个字符
paragraph_format = style.paragraph_format
paragraph_format.first_line_indent = Cm(0.74)
# 段落
p1 = document.add_paragraph('你们平时')
# 字体加粗
p1.add_run('Word文档').bold = True
# 斜体
p1.add_run('用的多吗？').italic = True
# 列表
document.add_paragraph('A：我们用的多', style='List Bullet')
document.add_paragraph('B：我们用的少', style='List Bullet')
document.add_paragraph('C：我们用的不多不少', style='List Bullet')
document.add_heading('标题2', 2)
# 段落
p2 = document.add_paragraph('我平时基本都是手动操作Word文档，现在打算利用Python来操作它，'
                       '你们平时是手动操作Word文档？如果是的话，')
run = p2.add_run('一起来了解下如何通过Python来操作吧！')
# 设置字体大小
run.font.size = Pt(12)
# 表格
table = document.add_table(rows=3, cols=2, style='Table Grid')
# 表头
hc = table.rows[0].cells
hc[0].text = '姓名'
hc[1].text = '年龄'
# 表体
bc1 = table.rows[1].cells
bc1[0].text = '张三'
bc1[1].text = '22'
bc2 = table.rows[2].cells
bc2[0].text = '李四'
bc2[1].text = '33'
# 分页
# document.add_page_break()
# 图片
document.add_picture(r'E:\仓库\图片\证件照.jpg', width=Inches(1))
# 保存
document.save('test.docx')'''


from docx import Document

# 打开文档
document = Document('test.docx')
# 读取标题、段落、列表内容
ps = [ paragraph.text for paragraph in document.paragraphs]
for p in ps:
    print(p)
# 读取表格内容
ts = [table for table in document.tables]
for t in ts:
    for row in t.rows:
        for cell in row.cells:
            print(cell.text, end=' ')
        print()