# Python 进阶（七）： Word 基本操作

# 1. 概述
# Word 是一个十分常用的文字处理工具，通常我们都是手动来操作它，本节我们来看一下如何通过 Python 来操作。
#
# Python 提供了 python-docx 库，该库就是为 Word 文档量身定制的，安装使用 pip install python-docx 命令即可。
#
# 2. 写入
# 首先，我们使用 Python 来创建一个 Word 文档并向其中写入一些内容。
#
# 2.1 标题
# 我们先来创建 Word 文档并向其中添加标题，完整实现代码如下所示：
'''from docx import Document

#创建文档
document =Document()
#标题
document.add_heading('标题0',0)#注释掉没有就没有标题了，应该是按照012来区别的
document.add_heading('标题1',1)
document.add_heading('标题2',2)
#保存
document.save('test.docx')'''

# 2.2 段落
# 我们接着向 Word 文档中添加段落内容，完整实现代码如下所示：

'''from docx import Document

#创建文档
document = Document()
#标题
document.add_heading('标题0',0)
document.add_heading('标题1',1)

#段落
document.add_paragraph('你们平时word文档用的多么？')
#列表
document.add_paragraph('A:我们用的多',style='List Bullet')
document.add_paragraph('B:我们用的少',style='List Bullet')
document.add_paragraph('C:我们用的不多不少',style='List Bullet')
document.add_heading('标题2',3)

#段落
document.add_paragraph('我们平时基本上都是手动操作word文档，现在打算利用python来操作它，'
                       '你们平时手动操作word1文档？如果是的haul，一起来了解下如何通过'
                       'python来操作吧！')
#保存
document.save('test.docx')
'''


# 2.3 表格
# 我们接着向文档中插入表格，完整实现代码如下所示：
'''from docx import Document

#创建文档
document = Document()
#标题
document.add_heading('标题0',0)
document.add_heading('标题1',1)

#段落
document.add_paragraph('你们平时word文档用的多么？')
#列表
document.add_paragraph('A:我们用的多',style='List Bullet')
document.add_paragraph('B:我们用的少',style='List Bullet')
document.add_paragraph('C:我们用的不多不少',style='List Bullet')
document.add_heading('标题2',3)
#段落
document.add_paragraph('我们平时基本上都是手动操作word文档，现在打算利用python来操作它，'
                       '你们平时手动操作word1文档？如果是的haul，一起来了解下如何通过'
                       'python来操作吧！')
#表格
table = document.add_table(rows=3,cols=2,style='Table Grid')
#表头
hc = table.rows[0].cells
hc[0].text = '姓名'
hc[1].text ='年龄'
#表体
bc1 = table.rows[1].cells
bc1[0].text = '张三'
bc1[1].text = '22'
bc2 = table.rows[2].cells
bc2[0].text = '李四'
bc2[1].text = '33'
#保存
document.save('text.docx')'''


# 2.4 图片
# 我们接着向文档中插入图片，完整实现代码如下所示：
'''from docx import Document
from docx.shared import Inches

#创建文档
document = Document()
#标题
document.add_heading('标题0',0)
document.add_heading('标题1',1)

#段落
document.add_paragraph('你们平时word文档用的多么？')
#列表
document.add_paragraph('A:我们用的多',style='List Bullet')
document.add_paragraph('B:我们用的少',style='List Bullet')
document.add_paragraph('C:我们用的不多不少',style='List Bullet')
document.add_heading('标题2',3)
#段落
document.add_paragraph('我们平时基本上都是手动操作word文档，现在打算利用python来操作它，'
                       '你们平时手动操作word1文档？如果是的haul，一起来了解下如何通过'
                       'python来操作吧！')
#表格
table = document.add_table(rows=3,cols=2,style='Table Grid')
#表头
hc = table.rows[0].cells
hc[0].text = '姓名'
hc[1].text ='年龄'
#表体
bc1 = table.rows[1].cells
bc1[0].text = '张三'
bc1[1].text = '22'
bc2 = table.rows[2].cells
bc2[0].text = '李四'
bc2[1].text = '33'
#分页
document.add_page_break()
#图片
document.add_picture(r'E:\仓库\图片\证件照.jpg',width=Inches(3),height=Inches(3))
#保存
document.save('text.docx')
'''

# 2.5 样式
# 我们再设置一下基本样式，比如：标题居中、字体加粗、首行缩进等，完整实现代码如下所示：

'''from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt

#创建文档
document = Document()
# style = document.styles('Normal')  TypeError: 'Styles' object is not callable  粗心大意方括号
style = document.styles['Normal']

#标题
t0 = document.add_heading('标题0', 0)
#居中
t0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
document.add_heading('标题1', 1)
#首行缩进两个字符
paragraph_format = style.paragraph_format
paragraph_format.first_line_indent = Cm(0.74)#indent单词拼写错误
#段落
p1 = document.add_paragraph('你们平时')
#字体加粗
p1.add_run('Word文档').bold = True
#斜体
p1.add_run('用的多么？').italic = True

#列表
document.add_paragraph('A:我们用的多', style='List Bullet')
document.add_paragraph('B:我们用的少', style='List Bullet')
document.add_paragraph('C:我们用的不多不少', style='List Bullet')
document.add_heading('标题2', 3)
#段落
p2 = document.add_paragraph('我们平时基本上都是手动操作word文档，现在打算利用python来操作它，'
                       '你们平时手动操作word1文档？如果是的haul，'
                      )
run = p2.add_run('一起来了解下如何通过python来操作吧！')
#设置字体大小
run.font.size = Pt(12)
#表格
table = document.add_table(rows=3,cols=2,style='Table Grid')
#表头
hc = table.rows[0].cells
hc[0].text = '姓名'
hc[1].text = '年龄'
#表体
bc1 = table.rows[1].cells
bc1[0].text = '张三'
bc1[1].text = '22'
bc2 = table.rows[2].cells
bc2[0].text = '李四'
bc2[1].text = '33'
#分页
document.add_page_break()
#图片
document.add_picture(r'E:\仓库\图片\证件照.jpg',width=Inches(3),height=Inches(3))
#保存
document.save('text.docx')'''
# 3. 读取
# 我们再来读取一下之前 Word 文档中写入的内容，完整代码实现如下所示：

from docx import  Document
#打开文档
document = Document('text.docx')
#读取标题、段落、列表内容
ps = [ paragraph.text for paragraph in document.paragraphs]#paragraph的r后面容易少一个a啊
for p in  ps:
    print(p)

#读取表格内容
ts = [table for table in document.tables]
for t in ts:
    for row in t.rows:
        for cell in row.cells:
            print(cell.text,end= " ")
        print()


